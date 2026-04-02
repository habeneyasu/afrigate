"""
Document Intelligence Agent
Extracts trade-related fields:
Plain text strings (typed input)
PDF files
DOCX files
TXT files
"""

from __future__ import annotations

import io
import re
import json
import string
from pathlib import Path
from typing import Any, Mapping
from dataclasses import dataclass, field, asdict


try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from rapidfuzz import process as fuzz_process, fuzz
    HAS_FUZZ = True
except ImportError:
    HAS_FUZZ = False




@dataclass
class TradeDocument:
    
    exporter_name: str | None = None
    exporter_address: str | None = None
    exporter_country: str | None = None
    exporter_tax_id: str | None = None
    exporter_email: str | None = None
    exporter_phone: str | None = None

    importer_name: str | None = None
    importer_address: str | None = None
    importer_country: str | None = None
    importer_tax_id: str | None = None
    importer_email: str | None = None
    importer_phone: str | None = None

    consignee_name: str | None = None
    consignee_address: str | None = None
    notify_party: str | None = None
    manufacturer_name: str | None = None
    manufacturer_country: str | None = None
    buyer_name: str | None = None
    seller_name: str | None = None
    freight_forwarder: str | None = None
    customs_broker: str | None = None
    shipping_agent: str | None = None

    product_name: str | None = None
    product_description: str | None = None
    product_category: str | None = None
    hs_code: str | None = None                    
    commodity_code: str | None = None
    tariff_code: str | None = None
    schedule_b_number: str | None = None
    eccn: str | None = None                      
    naics_code: str | None = None
    sic_code: str | None = None
    brand: str | None = None
    model_number: str | None = None
    batch_lot_number: str | None = None
    serial_number: str | None = None
    quantity: str | None = None
    quantity_unit: str | None = None
    net_weight: str | None = None
    gross_weight: str | None = None
    weight_unit: str | None = None
    volume: str | None = None
    volume_unit: str | None = None
    number_of_packages: str | None = None
    package_type: str | None = None
    dimensions: str | None = None
    country_of_origin: str | None = None
    country_of_manufacture: str | None = None

    invoice_value: str | None = None
    currency: str | None = None
    unit_price: str | None = None
    total_value: str | None = None
    cif_value: str | None = None                  
    fob_value: str | None = None                
    cfr_value: str | None = None        
    ex_works_value: str | None = None
    insurance_value: str | None = None
    freight_value: str | None = None
    customs_value: str | None = None
    duty_amount: str | None = None
    duty_rate: str | None = None
    vat_amount: str | None = None
    tax_amount: str | None = None
    discount: str | None = None
    incoterms: str | None = None                
    payment_terms: str | None = None
    letter_of_credit_number: str | None = None
    bank_name: str | None = None
    port_of_loading: str | None = None
    port_of_discharge: str | None = None
    port_of_entry: str | None = None
    final_destination: str | None = None
    origin_country: str | None = None
    destination_country: str | None = None
    transit_country: str | None = None
    mode_of_transport: str | None = None      
    vessel_name: str | None = None
    vessel_flag: str | None = None
    voyage_number: str | None = None
    flight_number: str | None = None
    truck_number: str | None = None
    train_number: str | None = None
    container_number: str | None = None
    container_type: str | None = None    
    seal_number: str | None = None
    bill_of_lading_number: str | None = None
    airway_bill_number: str | None = None
    master_bill_number: str | None = None
    house_bill_number: str | None = None
    booking_number: str | None = None
    shipping_line: str | None = None

    invoice_number: str | None = None
    invoice_date: str | None = None
    purchase_order_number: str | None = None
    contract_number: str | None = None
    customs_declaration_number: str | None = None
    entry_number: str | None = None
    export_license_number: str | None = None
    import_license_number: str | None = None
    certificate_of_origin_number: str | None = None
    phytosanitary_certificate_number: str | None = None
    health_certificate_number: str | None = None
    fumigation_certificate_number: str | None = None
    inspection_certificate_number: str | None = None
    conformity_certificate_number: str | None = None
    packing_list_number: str | None = None
    dangerous_goods_declaration: str | None = None
    msds_number: str | None = None            
    cites_permit_number: str | None = None      
    halal_certificate_number: str | None = None
    kosher_certificate_number: str | None = None
    organic_certificate_number: str | None = None
    iso_certificate_number: str | None = None
    gsp_form_number: str | None = None           
    eur1_number: str | None = None       
    form_a_number: str | None = None

    shipment_date: str | None = None
    departure_date: str | None = None
    arrival_date: str | None = None
    delivery_date: str | None = None
    expiry_date: str | None = None
    production_date: str | None = None
    best_before_date: str | None = None

    import_regime: str | None = None
    export_regime: str | None = None
    customs_procedure_code: str | None = None
    statistical_value: str | None = None
    preference_code: str | None = None
    quota_number: str | None = None
    anti_dumping_case: str | None = None
    sanctions_screening: str | None = None
    dual_use_flag: str | None = None
    end_use_certificate: str | None = None
    embargo_check: str | None = None
    free_trade_agreement: str | None = None      

    trade_type: str | None = None               
    transaction_type: str | None = None        
    reason_for_export: str | None = None
    end_use: str | None = None

    raw_text: str | None = None
    source_file: str | None = None
    source_format: str | None = None
    extraction_confidence: dict[str, str] = field(default_factory=dict)
    unmatched_lines: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        d = asdict(self)
        # remove internal meta from main output
        return {k: v for k, v in d.items() if v is not None and k not in
                ("raw_text", "unmatched_lines", "extraction_confidence")}

    def summary(self) -> str:
        filled = {k: v for k, v in self.to_dict().items()}
        lines = [f"  {k:<40} {v}" for k, v in filled.items()]
        return "Extracted Trade Fields\n" + "=" * 60 + "\n" + "\n".join(lines)


# label aliases

#maps canonical field name
FIELD_ALIASES: dict[str, list[str]] = {
    
    "exporter_name": [
        "exporter", "shipper", "sender", "shipped by", "exported by",
        "exporter name", "shipper name", "seller", "vendor",
    ],
    "exporter_address": [
        "exporter address", "shipper address", "exporter's address",
        "address of exporter",
    ],
    "exporter_country": ["exporter country", "shipper country", "country of exporter"],
    "exporter_tax_id": [
        "exporter vat", "exporter tax id", "tin", "vat number", "tax id",
        "gst number", "company reg", "registration number",
    ],
    "exporter_email": ["exporter email", "shipper email"],
    "exporter_phone": ["exporter phone", "shipper phone", "exporter tel", "shipper tel"],

    "importer_name": [
        "importer", "buyer", "consignee", "recipient", "importer name",
        "consignee name", "buyer name", "imported by",
    ],
    "importer_address": [
        "importer address", "buyer address", "consignee address",
        "address of importer",
    ],
    "importer_country": ["importer country", "buyer country", "destination country"],
    "importer_tax_id": ["importer vat", "importer tax id", "buyer tax id"],
    "importer_email": ["importer email", "buyer email"],
    "importer_phone": ["importer phone", "buyer phone"],

    "consignee_name": ["consignee", "consignee name", "deliver to"],
    "consignee_address": ["consignee address"],
    "notify_party": ["notify party", "notify", "also notify", "notification party"],
    "manufacturer_name": [
        "manufacturer", "made by", "manufactured by", "producer", "fabricator",
    ],
    "manufacturer_country": ["manufacturer country", "country of manufacture"],
    "freight_forwarder": [
        "freight forwarder", "forwarder", "forwarding agent", "logistics agent",
    ],
    "customs_broker": ["customs broker", "broker", "clearing agent"],
    "shipping_agent": ["shipping agent", "agent", "carrier agent"],

    #product
    "product_name": [
        "product", "goods", "commodity", "item", "description of goods",
        "goods description", "article", "merchandise", "cargo",
        "product name", "item name",
    ],
    "product_description": [
        "description", "full description", "detailed description",
        "description of commodities", "nature of goods",
    ],
    "product_category": ["category", "product category", "goods category", "type of goods"],
    "hs_code": [
        "hs code", "hs", "harmonized code", "tariff heading", "hs tariff",
        "commodity code", "tariff code", "customs tariff number",
        "hs number", "hts code", "hts", "schedule b",
    ],
    "commodity_code": ["commodity code", "commodity"],
    "tariff_code": ["tariff code", "tariff"],
    "eccn": ["eccn", "export control classification", "export control number"],
    "brand": ["brand", "brand name", "trademark", "trade name", "make"],
    "model_number": ["model", "model number", "model no", "part number", "part no"],
    "batch_lot_number": ["batch", "lot", "batch number", "lot number", "batch no", "lot no"],
    "serial_number": ["serial number", "serial no", "s/n"],
    "quantity": [
        "quantity", "qty", "amount", "units", "no of units", "number of units",
        "pieces", "pcs", "cartons", "bags", "bales", "rolls",
    ],
    "quantity_unit": ["unit", "uom", "unit of measure", "unit of measurement"],
    "net_weight": [
        "net weight", "net wt", "nw", "net", "net weight kg",
        "net weight lbs",
    ],
    "gross_weight": [
        "gross weight", "gross wt", "gw", "gross", "gross weight kg",
        "total weight",
    ],
    "weight_unit": ["weight unit", "wt unit"],
    "volume": ["volume", "cbm", "cubic meters", "measurement"],
    "volume_unit": ["volume unit"],
    "number_of_packages": [
        "packages", "no of packages", "number of packages", "pkgs",
        "no of pieces", "number of pieces", "total packages",
    ],
    "package_type": ["package type", "packing type", "type of package"],
    "dimensions": ["dimensions", "size", "l x w x h", "length x width x height"],
    "country_of_origin": [
        "country of origin", "origin", "made in", "manufactured in",
        "coo", "place of origin",
    ],
    "country_of_manufacture": [
        "country of manufacture", "manufacturing country",
    ],

    #value
    "invoice_value": [
        "invoice value", "total invoice", "invoice amount", "amount",
        "total amount", "value",
    ],
    "currency": ["currency", "ccy", "currency code"],
    "unit_price": ["unit price", "price per unit", "price", "rate", "unit cost"],
    "total_value": [
        "total value", "total", "grand total", "total price",
        "total amount due",
    ],
    "cif_value": ["cif", "cif value", "cost insurance freight"],
    "fob_value": ["fob", "fob value", "free on board"],
    "cfr_value": ["cfr", "c&f", "cnf", "cost and freight"],
    "ex_works_value": ["exw", "ex works", "ex-works", "ex factory"],
    "insurance_value": ["insurance", "insurance value", "insurance amount"],
    "freight_value": ["freight", "freight charges", "freight cost", "freight amount"],
    "customs_value": ["customs value", "dutiable value"],
    "duty_amount": ["duty", "duty amount", "customs duty", "import duty"],
    "duty_rate": ["duty rate", "tariff rate"],
    "vat_amount": ["vat", "vat amount", "value added tax"],
    "tax_amount": ["tax", "tax amount", "taxes"],
    "discount": ["discount", "rebate", "deduction"],
    "incoterms": [
        "incoterms", "incoterm", "terms of delivery", "delivery terms",
        "trade terms",
    ],
    "payment_terms": [
        "payment terms", "terms of payment", "payment conditions",
        "payment method",
    ],
    "letter_of_credit_number": ["letter of credit", "l/c", "lc number", "lc no"],
    "bank_name": ["bank", "bank name", "issuing bank", "remitting bank"],

    #shipment
    "port_of_loading": [
        "port of loading", "pol", "loading port", "port of departure",
        "from port", "origin port", "place of loading",
    ],
    "port_of_discharge": [
        "port of discharge", "pod", "discharge port", "port of arrival",
        "destination port", "unloading port", "place of discharge",
    ],
    "port_of_entry": ["port of entry", "entry port"],
    "final_destination": [
        "final destination", "place of delivery", "ultimate destination",
        "final delivery",
    ],
    "origin_country": [
        "origin country", "country of origin", "from country", "exporting country",
    ],
    "destination_country": [
        "destination country", "to country", "importing country",
        "country of destination",
    ],
    "transit_country": ["transit country", "country of transit", "transit via"],
    "mode_of_transport": [
        "mode of transport", "transport mode", "means of transport",
        "mode of shipment", "shipment mode",
    ],
    "vessel_name": ["vessel", "vessel name", "ship name", "mother vessel"],
    "vessel_flag": ["flag", "vessel flag", "flag of vessel"],
    "voyage_number": ["voyage", "voyage number", "voyage no", "voy no"],
    "flight_number": ["flight", "flight number", "flight no"],
    "truck_number": ["truck", "truck number", "truck no", "vehicle number"],
    "train_number": ["train", "train number"],
    "container_number": [
        "container", "container number", "container no", "cntr no", "cntr",
    ],
    "container_type": ["container type", "container size"],
    "seal_number": ["seal", "seal number", "seal no"],
    "bill_of_lading_number": [
        "bill of lading", "b/l", "bl number", "bl no", "b/l number",
        "bol", "lading number",
    ],
    "airway_bill_number": ["airway bill", "awb", "awb number", "air waybill"],
    "master_bill_number": ["master bill", "mbl", "master b/l"],
    "house_bill_number": ["house bill", "hbl", "house b/l"],
    "booking_number": ["booking", "booking number", "booking ref", "booking reference"],
    "shipping_line": ["shipping line", "carrier", "ocean carrier", "airline"],

    #docs
    "invoice_number": [
        "invoice number", "invoice no", "inv no", "inv number", "invoice #",
        "commercial invoice number",
    ],
    "invoice_date": ["invoice date", "date of invoice", "inv date"],
    "purchase_order_number": ["purchase order", "po number", "po no", "p.o.", "order number"],
    "contract_number": ["contract", "contract number", "contract no"],
    "customs_declaration_number": [
        "customs declaration", "declaration number", "sad number",
        "customs entry number",
    ],
    "export_license_number": [
        "export license", "export licence", "export permit",
        "export authorization",
    ],
    "import_license_number": [
        "import license", "import licence", "import permit",
        "import authorization",
    ],
    "certificate_of_origin_number": [
        "certificate of origin", "co number", "c/o number",
        "origin certificate",
    ],
    "phytosanitary_certificate_number": [
        "phytosanitary", "phyto certificate", "plant health certificate",
        "phytosanitary certificate",
    ],
    "health_certificate_number": [
        "health certificate", "sanitary certificate", "veterinary certificate",
        "animal health certificate",
    ],
    "fumigation_certificate_number": [
        "fumigation", "fumigation certificate", "treatment certificate",
    ],
    "inspection_certificate_number": [
        "inspection certificate", "inspection report", "survey report",
        "quality certificate",
    ],
    "conformity_certificate_number": [
        "certificate of conformity", "conformity certificate", "coc",
        "coc number",
    ],
    "packing_list_number": ["packing list", "packing list number", "pl number"],
    "dangerous_goods_declaration": [
        "dangerous goods", "dg declaration", "hazmat declaration",
        "imdg declaration",
    ],
    "msds_number": ["msds", "sds", "safety data sheet", "material safety data sheet"],
    "halal_certificate_number": ["halal", "halal certificate"],
    "kosher_certificate_number": ["kosher", "kosher certificate"],
    "organic_certificate_number": ["organic", "organic certificate", "organic certification"],
    "iso_certificate_number": ["iso", "iso certificate", "iso certification"],
    "gsp_form_number": ["gsp", "form a", "generalised system of preferences"],
    "eur1_number": ["eur1", "eur.1", "movement certificate"],
    "cites_permit_number": ["cites", "cites permit", "wildlife permit"],
    "free_trade_agreement": [
        "free trade agreement", "fta", "preferential agreement", "comesa",
        "afcfta", "ecowas", "agoa", "euaa",
    ],

    #dates
    "shipment_date": ["shipment date", "date of shipment", "ship date"],
    "departure_date": ["departure date", "date of departure", "etd"],
    "arrival_date": ["arrival date", "date of arrival", "eta"],
    "delivery_date": ["delivery date", "date of delivery", "expected delivery"],
    "expiry_date": ["expiry date", "expiration date", "valid until", "valid to"],
    "production_date": ["production date", "manufacture date", "date of production"],
    "best_before_date": ["best before", "use by", "best before date"],

    #compliance
    "import_regime": ["import regime", "customs regime", "import procedure"],
    "export_regime": ["export regime", "export procedure"],
    "customs_procedure_code": ["customs procedure", "procedure code", "cpc"],
    "preference_code": ["preference code", "tariff preference"],
    "quota_number": ["quota", "quota number", "quota reference"],
    "trade_type": [
        "trade type", "type of trade", "shipment type",
        "export", "import", "re-export", "transit shipment",
    ],
    "transaction_type": ["transaction type", "nature of transaction"],
    "reason_for_export": ["reason for export", "purpose of export", "reason"],
    "end_use": ["end use", "end use certificate", "intended use"],
}

#country lookup

KNOWN_COUNTRIES = [
    "Afghanistan", "Albania", "Algeria", "Angola", "Argentina", "Armenia",
    "Australia", "Austria", "Azerbaijan", "Bahrain", "Bangladesh", "Belgium",
    "Benin", "Bolivia", "Bosnia", "Botswana", "Brazil", "Bulgaria", "Burkina Faso",
    "Burundi", "Cambodia", "Cameroon", "Canada", "Chad", "Chile", "China",
    "Colombia", "Congo", "Costa Rica", "Croatia", "Cuba", "Cyprus",
    "Czech Republic", "Denmark", "Djibouti", "Dominican Republic", "DR Congo",
    "Ecuador", "Egypt", "El Salvador", "Eritrea", "Estonia", "Eswatini",
    "Ethiopia", "Finland", "France", "Gabon", "Gambia", "Georgia", "Germany",
    "Ghana", "Greece", "Guatemala", "Guinea", "Honduras", "Hungary", "India",
    "Indonesia", "Iran", "Iraq", "Ireland", "Israel", "Italy", "Ivory Coast",
    "Jamaica", "Japan", "Jordan", "Kazakhstan", "Kenya", "Kuwait", "Laos",
    "Latvia", "Lebanon", "Lesotho", "Liberia", "Libya", "Lithuania",
    "Luxembourg", "Madagascar", "Malawi", "Malaysia", "Mali", "Malta",
    "Mauritania", "Mauritius", "Mexico", "Moldova", "Morocco", "Mozambique",
    "Myanmar", "Namibia", "Nepal", "Netherlands", "New Zealand", "Nicaragua",
    "Niger", "Nigeria", "North Korea", "Norway", "Oman", "Pakistan",
    "Palestine", "Panama", "Paraguay", "Peru", "Philippines", "Poland",
    "Portugal", "Qatar", "Romania", "Russia", "Rwanda", "Saudi Arabia",
    "Senegal", "Serbia", "Sierra Leone", "Singapore", "Slovakia", "Slovenia",
    "Somalia", "South Africa", "South Korea", "South Sudan", "Spain",
    "Sri Lanka", "Sudan", "Sweden", "Switzerland", "Syria", "Taiwan",
    "Tanzania", "Thailand", "Togo", "Tunisia", "Turkey", "Uganda", "Ukraine",
    "United Arab Emirates", "United Kingdom", "United States", "Uruguay",
    "Uzbekistan", "Venezuela", "Vietnam", "Yemen", "Zambia", "Zimbabwe",
]

CURRENCY_CODES = {
    "USD", "EUR", "GBP", "JPY", "CNY", "CHF", "CAD", "AUD", "HKD", "SGD",
    "KES", "ETB", "NGN", "GHS", "ZAR", "EGP", "MAD", "TZS", "UGX", "RWF",
    "XOF", "XAF", "AED", "SAR", "INR", "BRL", "MXN", "RUB", "TRY", "SEK",
    "NOK", "DKK", "NZD", "THB", "IDR", "MYR", "PHP", "VND", "PKR", "BDT",
}

CURRENCY_SYMBOLS = {"$": "USD", "€": "EUR", "£": "GBP", "¥": "JPY", "₦": "NGN",
                    "₹": "INR", "R": "ZAR", "Br": "ETB"}

INCOTERM_CODES = {
    "EXW", "FCA", "FAS", "FOB", "CFR", "CIF", "CPT", "CIP",
    "DAP", "DPU", "DDP", "DAT", "DDU", "C&F", "CNF",
}

TRANSPORT_MODES = {
    "sea": ["sea", "ocean", "maritime", "vessel", "ship", "marine"],
    "air": ["air", "airfreight", "air freight", "airplane", "aircraft", "flight"],
    "road": ["road", "truck", "lorry", "vehicle", "highway", "land"],
    "rail": ["rail", "railway", "train", "railroad"],
    "multimodal": ["multimodal", "combined", "intermodal"],
}


# text extractions

def extract_text_from_pdf(path: str | Path) -> str:
    if not HAS_PDF:
        raise ImportError("pdfplumber is required: pip install pdfplumber")
    text_parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def extract_text_from_docx(path: str | Path) -> str:
    if not HAS_DOCX:
        raise ImportError("python-docx is required: pip install python-docx")
    doc = DocxDocument(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # also grab table cells
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))
    return "\n".join(parts)


def extract_text_from_txt(path: str | Path) -> str:
    return Path(path).read_text(encoding="utf-8", errors="replace")


def get_text(source: str | Path | bytes, fmt: str | None = None) -> tuple[str, str]:
    
    #returns (raw_text, detected_format)
    if isinstance(source, bytes):
        #detect from magic bytes
        if source[:4] == b"%PDF":
            path = Path("/tmp/_afrigate_upload.pdf")
            path.write_bytes(source)
            return extract_text_from_pdf(path), "pdf"
        elif source[:2] == b"PK":       
            path = Path("/tmp/_afrigate_upload.docx")
            path.write_bytes(source)
            return extract_text_from_docx(path), "docx"
        else:
            return source.decode("utf-8", errors="replace"), "txt"

    path = Path(source) if not isinstance(source, str) else None

    #plain string
    if isinstance(source, str) and (len(source) > 300 or "\n" in source
                                    or not Path(source).exists()):
        return source, "text"

    if path is None:
        path = Path(source)

    if not path.exists():
        
        return source, "text"

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path), "pdf"
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(path), "docx"
    elif suffix in (".txt", ".text", ".csv", ".log"):
        return extract_text_from_txt(path), "txt"
    else:
        
        try:
            return path.read_text(encoding="utf-8", errors="replace"), "txt"
        except Exception:
            return "", "unknown"



def _normalize(s: str) -> str:
    """Lowercase, collapse whitespace, strip punctuation on edges."""
    s = s.lower().strip()
    s = re.sub(r"\s+", " ", s)
    s = s.strip(string.punctuation + " ")
    return s


def _fuzzy_match_label(label: str, candidates: list[str],
                       threshold: int = 82) -> tuple[str | None, int]:
    #return (best_match, score) or (None, 0)
    if not HAS_FUZZ:
        return None, 0
    result = fuzz_process.extractOne(label, candidates, scorer=fuzz.ratio)
    if result and result[1] >= threshold:
        return result[0], result[1]
    return None, 0


#a reverse lookup: alias_normalized -> field_name
_ALIAS_LOOKUP: dict[str, str] = {}
_ALL_ALIASES: list[str] = []

for _field, _aliases in FIELD_ALIASES.items():
    for _alias in _aliases:
        _key = _normalize(_alias)
        _ALIAS_LOOKUP[_key] = _field
        _ALL_ALIASES.append(_key)


def _match_field_from_label(label: str) -> str | None:
    """Map a raw label string to a canonical field name."""
    norm = _normalize(label)

    #exact match
    if norm in _ALIAS_LOOKUP:
        return _ALIAS_LOOKUP[norm]

    #substring match
    for alias, field in _ALIAS_LOOKUP.items():
        if alias in norm or norm in alias:
            return field

    #fuzzy match
    match, score = _fuzzy_match_label(norm, _ALL_ALIASES)
    if match:
        return _ALIAS_LOOKUP.get(match)

    return None


#pattern extractors

_DATE_PATTERN = re.compile(
    r"\b(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}|\d{4}[\/\-\.]\d{2}[\/\-\.]\d{2}"
    r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{1,2},?\s*\d{4}"
    r"|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\b",
    re.IGNORECASE,
)

_MONEY_PATTERN = re.compile(
    r"(?P<symbol>[$€£¥₦₹R])?(?P<amount>\d[\d,]*(?:\.\d{1,4})?)\s*"
    r"(?P<code>USD|EUR|GBP|JPY|CNY|CHF|CAD|AUD|KES|ETB|NGN|GHS|ZAR|EGP|MAD|TZS|UGX|RWF)?",
    re.IGNORECASE,
)

_HS_PATTERN = re.compile(r"\b(\d{4}\.\d{2}(?:\.\d{2,4})?|\d{6,10})\b")

_WEIGHT_PATTERN = re.compile(
    r"(\d[\d,]*(?:\.\d+)?)\s*(kg|kgs|kilogram|kilograms|g|grams|lb|lbs|pound|pounds|"
    r"mt|mts|metric ton|metric tons|ton|tons|t)\b",
    re.IGNORECASE,
)

_QUANTITY_PATTERN = re.compile(
    r"(\d[\d,]*(?:\.\d+)?)\s*(pcs|pieces|units|bags|bales|cartons|boxes|rolls|"
    r"drums|pallets|cases|sheets|pairs|sets|litres|liters|l|ml|bottles)\b",
    re.IGNORECASE,
)

_CONTAINER_PATTERN = re.compile(
    r"\b([A-Z]{3}U\d{7}|[A-Z]{4}\d{6,7})\b"
)

_BL_PATTERN = re.compile(r"\b([A-Z]{2,6}\d{6,12})\b")

_EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")

_PHONE_PATTERN = re.compile(
    r"(?:\+\d{1,3}[\s\-]?)?\(?\d{1,4}\)?[\s\-]?\d{3,5}[\s\-]?\d{4,6}"
)


def _extract_incoterm(text: str) -> str | None:
    for code in INCOTERM_CODES:
        if re.search(r"\b" + re.escape(code) + r"\b", text, re.IGNORECASE):
            return code.upper()
    return None


def _extract_currency(text: str) -> str | None:
    # Currency code
    for code in CURRENCY_CODES:
        if re.search(r"\b" + code + r"\b", text, re.IGNORECASE):
            return code
    # Symbol
    for symbol, code in CURRENCY_SYMBOLS.items():
        if symbol in text:
            return code
    return None


def _extract_country(text: str) -> str | None:
    """Find a known country name in text (fuzzy)."""
    norm = text.strip()
    # Exact
    for c in KNOWN_COUNTRIES:
        if c.lower() in norm.lower():
            return c
    # Fuzzy
    if HAS_FUZZ:
        result = fuzz_process.extractOne(norm, KNOWN_COUNTRIES, scorer=fuzz.token_sort_ratio)
        if result and result[1] >= 80:
            return result[0]
    return norm if len(norm) > 1 else None


def _extract_transport_mode(text: str) -> str | None:
    norm = text.lower()
    for mode, keywords in TRANSPORT_MODES.items():
        for kw in keywords:
            if kw in norm:
                return mode
    return None


#Line-by-line key:value parser

_KV_PATTERNS = [
    # "Label: Value"  or  "Label – Value"  or  "Label | Value"
    re.compile(r"^(.+?)\s*[:–\-|]\s*(.+)$"),
    # "Label = Value"
    re.compile(r"^(.+?)\s*=\s*(.+)$"),
    # Table cell style  "Label    Value" (2+ spaces)
    re.compile(r"^(.+?)\s{2,}(.+)$"),
]


def _parse_kv_lines(text: str) -> list[tuple[str, str]]:
    """Extract (label, value) pairs from free text."""
    pairs: list[tuple[str, str]] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        for pat in _KV_PATTERNS:
            m = pat.match(line)
            if m:
                label, value = m.group(1).strip(), m.group(2).strip()
                if len(label) >= 2 and len(value) >= 1:
                    pairs.append((label, value))
                    break
    return pairs


#Sentence-level inline extraction

_INLINE_PATTERNS = [
    # "export X kg of Y from A to B"
    re.compile(
        r"export\s+(?P<quantity>[\d,]+(?:\.\d+)?)\s*(?P<qty_unit>\w+)?\s+"
        r"(?:of\s+)?(?P<product>.+?)\s+from\s+(?P<origin>.+?)\s+to\s+(?P<destination>.+)",
        re.IGNORECASE,
    ),
    # "ship X from A to B"
    re.compile(
        r"ship(?:ping|ment)?\s+(?P<product>.+?)\s+from\s+(?P<origin>.+?)\s+to\s+(?P<destination>.+)",
        re.IGNORECASE,
    ),
    # "import X from A"
    re.compile(
        r"import(?:ing|ation)?\s+(?P<product>.+?)\s+from\s+(?P<origin>.+)",
        re.IGNORECASE,
    ),
    # "value USD 8,000"
    re.compile(
        r"value\s+(?P<currency>[A-Z]{3})\s+(?P<amount>[\d,]+(?:\.\d+)?)",
        re.IGNORECASE,
    ),
    # "worth $8,000"
    re.compile(
        r"worth\s+(?P<symbol>[$€£₦])\s*(?P<amount>[\d,]+(?:\.\d+)?)",
        re.IGNORECASE,
    ),
]


def _apply_inline_patterns(text: str, doc: TradeDocument) -> None:
    for pat in _INLINE_PATTERNS:
        m = pat.search(text)
        if not m:
            continue
        gd = m.groupdict()

        if "product" in gd and gd["product"] and not doc.product_name:
            doc.product_name = gd["product"].strip().rstrip(".,;")
        if "quantity" in gd and gd["quantity"] and not doc.quantity:
            doc.quantity = gd["quantity"].replace(",", "")
        if "qty_unit" in gd and gd["qty_unit"] and not doc.quantity_unit:
            doc.quantity_unit = gd["qty_unit"]
        if "origin" in gd and gd["origin"] and not doc.origin_country:
            doc.origin_country = _extract_country(gd["origin"].strip().rstrip(".,;"))
            if not doc.country_of_origin:
                doc.country_of_origin = doc.origin_country
        if "destination" in gd and gd["destination"] and not doc.destination_country:
            doc.destination_country = _extract_country(gd["destination"].strip().rstrip(".,;"))
        if "currency" in gd and gd["currency"] and not doc.currency:
            doc.currency = gd["currency"].upper()
        if "amount" in gd and gd["amount"] and not doc.invoice_value:
            doc.invoice_value = gd["amount"].replace(",", "")
        if "symbol" in gd and gd["symbol"] and not doc.currency:
            doc.currency = CURRENCY_SYMBOLS.get(gd["symbol"], gd["symbol"])


#global pattern scan (whole text)

def _scan_global_patterns(text: str, doc: TradeDocument) -> None:
    """Scan entire text for patterns regardless of label context."""

    # HS Code
    if not doc.hs_code:
        m = _HS_PATTERN.search(text)
        if m:
            val = m.group(1)
            if len(val) >= 6:
                doc.hs_code = val

    # Incoterms
    if not doc.incoterms:
        doc.incoterms = _extract_incoterm(text)

    # Currency
    if not doc.currency:
        doc.currency = _extract_currency(text)

    # Email
    if not doc.exporter_email:
        m = _EMAIL_PATTERN.search(text)
        if m:
            doc.exporter_email = m.group(0)

    # Transport mode
    if not doc.mode_of_transport:
        doc.mode_of_transport = _extract_transport_mode(text)

    # Container
    if not doc.container_number:
        m = _CONTAINER_PATTERN.search(text)
        if m:
            doc.container_number = m.group(1)

    # Weight
    if not doc.net_weight:
        m = _WEIGHT_PATTERN.search(text)
        if m:
            doc.net_weight = m.group(1).replace(",", "")
            doc.weight_unit = m.group(2).lower()

    # Quantity
    if not doc.quantity:
        m = _QUANTITY_PATTERN.search(text)
        if m:
            doc.quantity = m.group(1).replace(",", "")
            if not doc.quantity_unit:
                doc.quantity_unit = m.group(2).lower()

    # Country of origin
    for phrase in ["made in", "origin:", "from"]:
        idx = text.lower().find(phrase)
        if idx != -1:
            snippet = text[idx + len(phrase): idx + len(phrase) + 40]
            snippet_clean = snippet.split("\n")[0].strip()
            # Skip cert-number-like values (e.g. CO-2024-ET-00712)
            if re.search(r'[A-Z]{2,}-\d{4}-', snippet_clean):
                continue
            country = _extract_country(snippet_clean)
            if country and not doc.country_of_origin:
                doc.country_of_origin = country
                break


#main

def extract_trade_fields(
    source: str | Path | bytes,
    source_name: str | None = None,
) -> TradeDocument:
    """
    entry point.

    Parameters
    ----------
    source   : str path, Path, raw bytes, or plain text string
    source_name : optional filename hint for format detection

    Returns
    -------
    TradeDocument with all detected fields populated.
    """
    raw_text, fmt = get_text(source)

    def _resolve_source_file():
        if source_name:
            return source_name
        if isinstance(source, bytes):
            return "<bytes>"
        s = str(source)
        if len(s) <= 260:
            try:
                if Path(s).exists():
                    return s
            except OSError:
                pass
        return "<inline>"

    doc = TradeDocument(
        raw_text=raw_text,
        source_file=_resolve_source_file(),
        source_format=fmt,
    )

    if not raw_text.strip():
        return doc

    #pass 1: Inline / sentence patterns
    _apply_inline_patterns(raw_text, doc)

    #pass 2: Key-Value line parsing
    kv_pairs = _parse_kv_lines(raw_text)
    unmatched = []

    for label, value in kv_pairs:
        field_name = _match_field_from_label(label)
        if field_name:
            current = getattr(doc, field_name, None)
            if current is None:
                # Post-process value by field type
                processed = _post_process_value(field_name, value, raw_text)
                setattr(doc, field_name, processed)
                doc.extraction_confidence[field_name] = "high"
        else:
            unmatched.append(f"{label}: {value}")

    doc.unmatched_lines = unmatched

    #pass 3: Global pattern scan 
    _scan_global_patterns(raw_text, doc)

    #pass 4: Cross-field inference
    _infer_fields(doc)

    return doc


def _post_process_value(field: str, value: str, full_text: str) -> str:
    """Clean and enrich a raw value string for a given field."""
    value = value.strip().strip('"\'')

    date_fields = {
        "invoice_date", "shipment_date", "departure_date", "arrival_date",
        "delivery_date", "expiry_date", "production_date", "best_before_date",
    }
    country_fields = {
        "exporter_country", "importer_country", "country_of_origin",
        "country_of_manufacture", "origin_country", "destination_country",
        "transit_country", "manufacturer_country",
    }
    money_fields = {
        "invoice_value", "total_value", "cif_value", "fob_value", "cfr_value",
        "unit_price", "freight_value", "insurance_value", "customs_value",
        "duty_amount", "vat_amount", "tax_amount", "ex_works_value",
    }

    if field in date_fields:
        m = _DATE_PATTERN.search(value)
        if m:
            return m.group(0)

    if field in country_fields:
        # Reject strings that look like reference/cert numbers (e.g. CO-2024-ET-00712)
        if re.search(r'[A-Z]{2,}-\d{4}-', value):
            return value
        return _extract_country(value) or value

    if field in money_fields:
        # Strip currency symbols and codes, keep the number
        cleaned = re.sub(r"[A-Z]{3}|[$€£¥₦₹]", "", value).strip().replace(",", "")
        return cleaned if re.search(r"\d", cleaned) else value

    if field == "incoterms":
        return _extract_incoterm(value) or value.upper()

    if field == "mode_of_transport":
        return _extract_transport_mode(value) or value

    if field in {"net_weight", "gross_weight"}:
        m = _WEIGHT_PATTERN.search(value)
        if m:
            return m.group(1).replace(",", "")

    return value


def _infer_fields(doc: TradeDocument) -> None:
    """Derive fields from each other where possible."""

    # If exporter_country missing, try country_of_origin
    if not doc.exporter_country and doc.country_of_origin:
        doc.exporter_country = doc.country_of_origin

    # Sync origin_country / country_of_origin
    if doc.origin_country and not doc.country_of_origin:
        doc.country_of_origin = doc.origin_country
    if doc.country_of_origin and not doc.origin_country:
        doc.origin_country = doc.country_of_origin

    # importer_country ↔ destination_country
    if doc.destination_country and not doc.importer_country:
        doc.importer_country = doc.destination_country
    if doc.importer_country and not doc.destination_country:
        doc.destination_country = doc.importer_country

    # importer_name ↔ consignee_name
    if doc.consignee_name and not doc.importer_name:
        doc.importer_name = doc.consignee_name

    # Infer trade type
    if not doc.trade_type:
        text = (doc.raw_text or "").lower()
        if "export" in text:
            doc.trade_type = "export"
        elif "import" in text:
            doc.trade_type = "import"
        elif "transit" in text:
            doc.trade_type = "transit"
        elif "re-export" in text or "re export" in text:
            doc.trade_type = "re-export"

    # Free trade agreement inference
    if not doc.free_trade_agreement:
        text = (doc.raw_text or "").lower()
        for fta in ["comesa", "afcfta", "ecowas", "agoa", "sadc", "igad",
                    "eac", "tripartite fta"]:
            if fta in text:
                doc.free_trade_agreement = fta.upper()
                break


# ---------------------------------------------------------------------------
# LangGraph node — maps TradeDocument → AfrigateState.extracted_fields
# ---------------------------------------------------------------------------

_COUNTRY_NAME_TO_ISO: dict[str, str] = {
    "ethiopia": "ET",
    "kenya": "KE",
    "ghana": "GH",
    "nigeria": "NG",
    "south africa": "ZA",
}


def _country_name_to_iso(name: str | None) -> str:
    """Map a country name or 2-letter code to ISO 3166-1 alpha-2."""
    if not name:
        return ""
    n = str(name).strip()
    if len(n) == 2 and n.isalpha():
        return n.upper()
    key = n.lower().strip().rstrip(".,;")
    return _COUNTRY_NAME_TO_ISO.get(key, "")


def _load_required_documents(destination_iso: str) -> list[str]:
    path = Path(__file__).resolve().parent.parent / "rules" / "regulations.json"
    data = json.loads(path.read_text(encoding="utf-8"))
    block = data.get(destination_iso) or {}
    return list(block.get("required_documents") or [])


def _documents_present_from_trade_doc(doc: TradeDocument) -> list[str]:
    """Infer structured compliance tokens from populated certificate / invoice fields."""
    present: list[str] = []
    if doc.certificate_of_origin_number:
        present.append("certificate_of_origin")
    if doc.phytosanitary_certificate_number:
        present.append("phytosanitary_certificate")
    if doc.invoice_number:
        present.append("commercial_invoice")
    if doc.packing_list_number:
        present.append("packing_list")
    return present


def _parse_weight_kg(doc: TradeDocument) -> float | None:
    unit_ok = {"kg", "kgs", "kilogram", "kilograms"}
    if doc.quantity and doc.quantity_unit and doc.quantity_unit.lower() in unit_ok:
        try:
            return float(str(doc.quantity).replace(",", ""))
        except ValueError:
            pass
    if doc.net_weight and doc.weight_unit and doc.weight_unit.lower() in unit_ok:
        try:
            return float(str(doc.net_weight).replace(",", ""))
        except ValueError:
            pass
    return None


def _parse_value_usd(doc: TradeDocument) -> float | None:
    for val_field in (doc.total_value, doc.invoice_value, doc.fob_value, doc.cif_value):
        if not val_field:
            continue
        try:
            return float(str(val_field).replace(",", ""))
        except ValueError:
            continue
    return None


def run_doc_intel(state: Mapping[str, Any]) -> dict[str, Any]:
    """Extract trade fields from ``document_raw`` for the LangGraph state.

    On iteration > 0 (after evaluator ``retry``), Phase 1 attaches all documents
    required by ``rules/regulations.json`` for the destination — mocked upload.
    """
    raw = state.get("document_raw") or ""
    doc = extract_from_text(str(raw))

    origin_iso = _country_name_to_iso(doc.origin_country or doc.country_of_origin)
    dest_iso = _country_name_to_iso(
        doc.destination_country or doc.importer_country
    )

    product = (doc.product_name or doc.product_description or "").strip().lower()

    iteration = int(state.get("iteration") or 0)
    documents_present = _documents_present_from_trade_doc(doc)
    if iteration > 0:
        required = _load_required_documents(dest_iso)
        documents_present = sorted(set(documents_present) | set(required))

    extracted_fields: dict[str, Any] = {
        "exporter": doc.exporter_name,
        "importer": doc.importer_name,
        "origin_country": origin_iso,
        "destination_country": dest_iso,
        "product": product,
        "weight_kg": _parse_weight_kg(doc),
        "value_usd": _parse_value_usd(doc),
        "documents_present": documents_present,
    }

    dest_disp = dest_iso or "?"
    log_line = (
        f"doc_intel: extracted product={product!r}, destination={dest_disp!r}"
    )
    return {"extracted_fields": extracted_fields, "agent_log": [log_line]}


#public api

def extract_from_text(text: str) -> TradeDocument:
    """Extract from a plain text string (typed or pasted)."""
    return extract_trade_fields(text, source_name="<typed input>")


def extract_from_file(file_path: str | Path) -> TradeDocument:
    """Extract from a PDF, DOCX, or TXT file on disk."""
    return extract_trade_fields(file_path, source_name=str(file_path))


def extract_from_bytes(data: bytes, filename: str = "") -> TradeDocument:
    """Extract from raw file bytes (e.g. Gradio upload)."""
    return extract_trade_fields(data, source_name=filename)


def to_json(doc: TradeDocument, indent: int = 2) -> str:
    return json.dumps(doc.to_dict(), indent=indent, ensure_ascii=False)


# cli

if __name__ == "__main__":
    import sys

    samples = [
        # ── Typed sentence ────────────────────────────────────────────────
        "Export 500kg roasted coffee from Ethiopia to Kenya, value USD 8,000.",

        # ── Structured key-value ──────────────────────────────────────────
        """
        Invoice Number: INV-2024-00712
        Invoice Date: 15 March 2024
        Exporter: Addis Trading PLC
        Exporter Country: Ethiopia
        Importer: Nairobi Imports Ltd
        Importer Country: Kenya
        Product: Roasted Arabica Coffee
        HS Code: 0901.21.00
        Quantity: 500 kg
        Net Weight: 500 kg
        Gross Weight: 540 kg
        Unit Price: USD 16.00
        Total Value: USD 8,000
        Currency: USD
        Incoterms: FOB
        Port of Loading: Djibouti Port
        Port of Discharge: Mombasa Port
        Mode of Transport: Sea
        Certificate of Origin: CO-2024-ET-00712
        Phytosanitary Certificate: PHY-2024-00123
        Free Trade Agreement: COMESA
        Payment Terms: 30 days LC
        """,
    ]

    for i, sample in enumerate(samples, 1):
        print(f"\n{'='*70}")
        print(f"SAMPLE {i}")
        print(f"{'='*70}")
        doc = extract_from_text(sample)
        print(doc.summary())
        if doc.unmatched_lines:
            print(f"\n  [Unmatched lines]: {doc.unmatched_lines}")

    # File path from CLI
    if len(sys.argv) > 1:
        fp = Path(sys.argv[1])
        print(f"\n{'='*70}")
        print(f"FILE: {fp}")
        print(f"{'='*70}")
        doc = extract_from_file(fp)
        print(doc.summary())
        print("\nJSON output:")
        print(to_json(doc))